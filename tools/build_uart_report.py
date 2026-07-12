from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "outputs" / "uart_fifo_verification" / "UART_FIFO_Verification_Report.docx"
ASSETS = ROOT / "report_assets"

NAVY = "002060"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
SOFT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN = "E2F0D9"
GREEN_TEXT = "006100"
MUTED = "667085"
WHITE = "FFFFFF"
BORDER = "B4C6E7"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")
    set_cell_margins(cell)


def set_table_geometry(table, widths, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = table._tbl.tblGrid
    new_grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        new_grid.append(col)
    table._tbl.replace(old_grid, new_grid)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color=BORDER, size=4):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def format_table_text(table, header=True, body_size=8.5):
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    if header and row_index == 0:
                        set_run_font(run, size=8.5, color=WHITE, bold=True)
                    else:
                        set_run_font(run, size=body_size, color="1F2937")
        if header and row_index == 0:
            for cell in row.cells:
                set_cell_shading(cell, NAVY)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if header:
        set_repeat_table_header(table.rows[0])
    set_table_borders(table)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        lead.bold = True
        paragraph.add_run(text[len(bold_lead):])
    else:
        paragraph.add_run(text)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(text, style="List Bullet")
    return paragraph


def add_numbered(doc, text):
    paragraph = doc.add_paragraph(text, style="List Number")
    return paragraph


def add_callout(doc, title, text, fill=CALLOUT, title_color=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color=BORDER, size=5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(title)
    set_run_font(run, size=10.5, color=title_color, bold=True)
    paragraph = cell.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run, size=10, color="1F2937")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_caption(doc, text):
    paragraph = doc.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, size=9, color=MUTED, italic=True)
    return paragraph


def add_code_listing(doc, file_path, heading, break_before=False):
    if break_before:
        doc.add_page_break()
    add_heading(doc, heading, 1)
    add_body(doc, f"Simulated source: {file_path.relative_to(ROOT).as_posix()}")
    lines = file_path.read_text(encoding="utf-8").splitlines()
    chunk_size = 44
    for start in range(0, len(lines), chunk_size):
        chunk = lines[start:start + chunk_size]
        paragraph = doc.add_paragraph(style="Code Block")
        paragraph.paragraph_format.keep_together = False
        paragraph.paragraph_format.page_break_before = False
        for line_index, line in enumerate(chunk):
            run = paragraph.add_run(line)
            if line_index != len(chunk) - 1:
                run.add_break(WD_BREAK.LINE)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("1F2937")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        1: (16, BLUE, 18, 10),
        2: (13, BLUE, 14, 7),
        3: (12, DARK_BLUE, 10, 5),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(6)

    code_style = styles.add_style("Code Block", 1)
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code_style.font.size = Pt(6.5)
    code_style.font.color.rgb = RGBColor.from_string("1F2937")
    code_style.paragraph_format.left_indent = Inches(0.08)
    code_style.paragraph_format.right_indent = Inches(0.08)
    code_style.paragraph_format.space_before = Pt(0)
    code_style.paragraph_format.space_after = Pt(4)
    code_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    code_style.paragraph_format.line_spacing = Pt(7.5)
    p_pr = code_style._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F7F8FA")
    p_pr.append(shd)


def build_report():
    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("UART VERIFICATION REPORT  |  AMD VIVADO 2024.1")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    add_page_number(section.footer.paragraphs[0])

    # Editorial-cover opening.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(56)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    run = kicker.add_run("DIGITAL DESIGN VERIFICATION REPORT")
    set_run_font(run, size=11, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("UART Transceiver with TX and RX FIFOs")
    set_run_font(run, size=28, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(30)
    run = subtitle.add_run("RTL implementation, self-checking test cases, Vivado XSim waveforms, and results")
    set_run_font(run, size=13.5, color=DARK_BLUE, italic=True)

    metadata = doc.add_table(rows=4, cols=2)
    values = [
        ("Protocol", "8-N-1 UART, LSB first, internal loopback"),
        ("Buffering", "16x8 synchronous FIFO on both TX and RX sides"),
        ("Verification", "12 simple self-checking test cases in AMD Vivado 2024.1 XSim"),
        ("Date", "11 July 2026"),
    ]
    for row, (label, value) in zip(metadata.rows, values):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], SOFT_BLUE)
        for run in row.cells[0].paragraphs[0].runs:
            set_run_font(run, size=10, color=NAVY, bold=True)
        for run in row.cells[1].paragraphs[0].runs:
            set_run_font(run, size=10, color="1F2937")
    set_table_geometry(metadata, [1800, 7560])
    format_table_text(metadata, header=False, body_size=10)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    add_callout(
        doc,
        "Final result",
        "PASS - all 12 test cases completed with zero errors. XSim reached the final PASS banner at 4098 ns.",
        fill=GREEN,
        title_color=GREEN_TEXT,
    )
    doc.add_page_break()

    add_heading(doc, "1. Executive Summary", 1)
    add_body(
        doc,
        "This report documents the completed UART loopback design after adding a receiver-side FIFO to the existing transmit-buffered architecture. The design now contains independent 16-byte FIFOs on both sides of the serial link. The TX FIFO absorbs fast producer bursts, while the RX FIFO holds completed bytes until the consumer reads them. The original external receive interface remains simple: rdy indicates that at least one byte is buffered, data_out shows the oldest byte, and rdy_clr pops one byte."
    )
    add_body(
        doc,
        "The verification environment deliberately uses basic, readable stimulus. It checks reset/idle behavior, ordinary and edge data patterns, back-to-back TX writes, delayed RX reads, ordered FIFO behavior, reset during activity, post-reset recovery, FIFO full protection, simultaneous FIFO access, and pointer wrap-around. The testbench is self-checking and reports failure with $fatal instead of relying only on visual waveform inspection."
    )

    summary_table = doc.add_table(rows=7, cols=2)
    summary_values = [
        ("Design under test", "uart_top with baud generator, TX FSM, RX FSM, TX FIFO, and RX FIFO"),
        ("Simulation top", "uart_top_tb"),
        ("Simulation parameters", "TX_DIV=16, RX_DIV=1 for accelerated 16x oversampling"),
        ("Test cases", "12"),
        ("Passed / failed", "12 / 0"),
        ("Simulation end", "4098 ns"),
        ("Acceptance result", "UART_FIFO_VERIFICATION_PASS"),
    ]
    for row, (label, value) in zip(summary_table.rows, summary_values):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], SOFT_BLUE)
        for run in row.cells[0].paragraphs[0].runs:
            set_run_font(run, size=9.5, color=NAVY, bold=True)
    set_table_geometry(summary_table, [2700, 6660])
    format_table_text(summary_table, header=False, body_size=9.5)

    add_heading(doc, "2. Design Architecture", 1)
    add_heading(doc, "2.1 End-to-end data flow", 2)
    add_body(doc, "The internal loopback path is:")
    flow = doc.add_paragraph(style="Code Block")
    flow.add_run(
        "data_in/wr_en -> TX FIFO -> UART transmitter -> tx_line -> UART receiver\n"
        "              -> completed rx_byte -> RX FIFO -> data_out/rdy/rdy_clr"
    )
    add_body(
        doc,
        "The baud-rate generator supplies one transmit enable pulse per serial bit and one receiver enable pulse per oversampling position. In the accelerated test configuration, the transmitter advances once every 16 clocks and the receiver advances every clock, preserving the intended 16:1 timing relationship."
    )

    add_heading(doc, "2.2 Transmit side", 2)
    add_bullet(doc, "A write pulse stores data_in in the TX FIFO if it is not full.")
    add_bullet(doc, "When the transmitter is idle and the FIFO is not empty, tx_start pops the oldest queued byte and starts a frame.")
    add_bullet(doc, "The sender emits a low start bit, eight LSB-first data bits, and a high stop bit.")
    add_bullet(doc, "busy remains high while the sender FSM is outside IDLE.")

    add_heading(doc, "2.3 Receiver-side FIFO integration", 2)
    add_body(
        doc,
        "The receiver still reconstructs one byte at a time, but its rdy output is now an internal completion indication called rx_byte_ready. When that indication is high and the RX FIFO has space, uart_top asserts rx_fifo_write. The same pulse acknowledges the receiver, ensuring that each completed byte is enqueued exactly once."
    )
    add_body(
        doc,
        "At the external interface, rdy is assigned to the inverse of rx_fifo_empty, data_out is assigned to the FIFO show-ahead output, and rdy_clr becomes the FIFO read enable. A consumer can therefore delay reading several bytes without losing order or overwriting the first unread byte."
    )
    add_callout(
        doc,
        "RX FIFO full behavior",
        "If the RX FIFO is full, rx_byte_ready remains asserted because the receiver is not acknowledged. Once a location becomes available, the pending byte is written and the receiver ready flag is cleared.",
        fill=CALLOUT,
    )

    add_heading(doc, "2.4 FIFO implementation", 2)
    add_body(
        doc,
        "uart_fifo is a 16-entry, 8-bit synchronous circular buffer with show-ahead output. Four-bit read and write pointers select the memory entries, while a five-bit count distinguishes empty from full even when pointer values are equal. Writes while full and reads while empty are ignored. During a simultaneous valid read and write, both pointers advance and the count remains unchanged."
    )

    add_heading(doc, "3. Verification Method", 1)
    add_heading(doc, "3.1 Self-checking structure", 2)
    add_numbered(doc, "Drive a simple stimulus using reusable tasks such as send_byte, pop_expect_byte, fifo_write, and fifo_pop_expect.")
    add_numbered(doc, "Wait for the exact handshake or FIFO state relevant to the case.")
    add_numbered(doc, "Compare the observed value with the expected value using case-aware checks where unknowns matter.")
    add_numbered(doc, "Increment error_count and print a FAIL message when any check does not match.")
    add_numbered(doc, "Print the final PASS banner only when error_count is zero; otherwise call $fatal.")
    add_body(
        doc,
        "A separate timeout process also calls $fatal if the design stops making progress. This prevents a hung wait statement from being mistaken for success."
    )

    add_heading(doc, "3.2 Vivado artifacts", 2)
    artifacts = doc.add_table(rows=5, cols=2)
    artifact_rows = [
        ("Project", "vivado/uart_fifo_verification/uart_fifo_verification.xpr"),
        ("Wave configuration", "vivado/uart_fifo_verification/uart_fifo_waveform.wcfg"),
        ("Exact waveform data", "vivado/uart_fifo_verification/uart_fifo_waveform.vcd"),
        ("XSim transcript", "vivado/.../xsim/vivado_xsim_results.log"),
        ("Testbench", "rtlCode/uart_top_tb.v"),
    ]
    for row, values in zip(artifacts.rows, artifact_rows):
        row.cells[0].text, row.cells[1].text = values
        set_cell_shading(row.cells[0], SOFT_BLUE)
        for run in row.cells[0].paragraphs[0].runs:
            set_run_font(run, size=9, color=NAVY, bold=True)
    set_table_geometry(artifacts, [2100, 7260])
    format_table_text(artifacts, header=False, body_size=9)

    add_heading(doc, "4. Test Cases and Results", 1)
    add_body(
        doc,
        "All cases below were executed in one repeatable XSim run. The end time is the timestamp printed by the testbench after the final assertion for that case passed."
    )
    test_rows = [
        ("TC01", "Reset / idle", "Release reset and inspect idle outputs.", "PASS at 10 ns"),
        ("TC02", "Single byte", "Send A5 and read it from RX FIFO.", "PASS at 350 ns"),
        ("TC03", "All zero", "Send 00.", "PASS at 670 ns"),
        ("TC04", "All one", "Send FF.", "PASS at 990 ns"),
        ("TC05", "Alternating", "Send 55 and AA.", "PASS at 1630 ns"),
        ("TC06", "TX FIFO burst", "Queue 11,22,33,44,77,88 without waiting for the UART.", "PASS at 3547 ns"),
        ("TC07", "RX FIFO hold/read", "Withhold reads, verify count/front, then pop six bytes in order.", "PASS at 3590 ns"),
        ("TC08", "Reset during transfer", "Queue C3/5A, reset while busy, verify flush and idle.", "PASS at 3606 ns"),
        ("TC09", "Post-reset recovery", "Send and receive 3C after reset.", "PASS at 3946 ns"),
        ("TC10", "FIFO capacity", "Fill 16 entries, reject EE overflow, read 0 through 15.", "PASS at 4078 ns"),
        ("TC11", "Simultaneous FIFO I/O", "Read 5C while writing D2; count remains one.", "PASS at 4090 ns"),
        ("TC12", "FIFO wrap-around reuse", "Write and read 9B after a complete pointer cycle.", "PASS at 4098 ns"),
    ]
    test_table = doc.add_table(rows=1, cols=4)
    for cell, text in zip(test_table.rows[0].cells, ("ID", "Area", "Simple stimulus", "Observed result")):
        cell.text = text
    for item in test_rows:
        row = test_table.add_row()
        for cell, text in zip(row.cells, item):
            cell.text = text
        set_cell_shading(row.cells[3], GREEN)
        for run in row.cells[3].paragraphs[0].runs:
            set_run_font(run, size=8.5, color=GREEN_TEXT, bold=True)
    set_table_geometry(test_table, [700, 1600, 4800, 2260])
    format_table_text(test_table, header=True, body_size=8.5)

    add_heading(doc, "4.1 Final XSim transcript", 2)
    transcript = [
        "TEST_CASE TC01 PASS reset and idle checks at 10 ns",
        "TEST_CASE TC02 PASS single byte A5 at 350 ns",
        "TEST_CASE TC03 PASS all-zero byte at 670 ns",
        "TEST_CASE TC04 PASS all-one byte at 990 ns",
        "TEST_CASE TC05 PASS alternating patterns at 1630 ns",
        "TEST_CASE TC06 PASS TX FIFO accepted six-byte burst at 3547 ns",
        "TEST_CASE TC07 PASS RX FIFO hold and ordered reads at 3590 ns",
        "TEST_CASE TC08 PASS reset flush during transfer at 3606 ns",
        "TEST_CASE TC09 PASS post-reset recovery at 3946 ns",
        "TEST_CASE TC10 PASS FIFO full and overflow protection at 4078 ns",
        "TEST_CASE TC11 PASS simultaneous FIFO read and write at 4090 ns",
        "TEST_CASE TC12 PASS FIFO wrap-around reuse at 4098 ns",
        "UART_FIFO_VERIFICATION_PASS: all 12 test cases passed at 4098 ns",
    ]
    paragraph = doc.add_paragraph(style="Code Block")
    for index, line in enumerate(transcript):
        run = paragraph.add_run(line)
        if index != len(transcript) - 1:
            run.add_break(WD_BREAK.LINE)

    add_heading(doc, "5. Exact Waveform Evidence", 1)
    add_body(
        doc,
        "The following figures are rendered from the VCD exported by AMD Vivado 2024.1 XSim. The VCD time resolution is 1 ps, so the transitions and bus values come from the simulator database rather than from a hand-drawn timing diagram."
    )

    add_caption(doc, "Figure 1. Complete 0-4100 ns UART and FIFO verification waveform")
    doc.add_picture(str(ASSETS / "uart_xsim_waveform_overview.png"), width=Inches(6.35))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_heading(doc, "5.1 Overview explanation", 2)
    add_bullet(doc, "Phases 2-5 show A5, 00, FF, 55, and AA passing through the complete UART path.")
    add_bullet(doc, "Each wr_en pulse places a byte into the TX path; busy stays high during its serial frame and tx_line shows the 8-N-1 bit sequence.")
    add_bullet(doc, "rx_byte_ready pulses once after each valid stop bit. For the single-byte cases, rdy_clr immediately pops the received byte.")
    add_bullet(doc, "During phase 6, TX FIFO count falls from five to zero as the six queued bytes are serialized, while RX FIFO count rises because reads are intentionally withheld.")
    add_bullet(doc, "The reset pulse near 3590 ns clears both FIFOs, lowers busy/rdy, and returns tx_line to its high idle state.")

    doc.add_page_break()
    add_caption(doc, "Figure 2. TX burst and RX FIFO accumulation/dequeue detail (TC06-TC07)")
    doc.add_picture(str(ASSETS / "uart_xsim_waveform_rx_fifo_detail.png"), width=Inches(6.35))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_heading(doc, "5.2 TX and RX FIFO explanation", 2)
    add_bullet(doc, "Six narrow wr_en pulses enqueue 11, 22, 33, 44, 77, and 88 much faster than the UART can transmit them.")
    add_bullet(doc, "TX FIFO count begins at five because the first byte is transferred to the sender while the remaining five wait in the queue.")
    add_bullet(doc, "Each completed frame creates one rx_byte_ready pulse and advances rx_byte through 11, 22, 33, 44, 77, and 88.")
    add_bullet(doc, "RX FIFO count climbs from zero to six while no consumer reads occur. rdy remains high and data_out remains 11, proving show-ahead behavior and stable oldest-byte ordering.")
    add_bullet(doc, "At the end of phase 7, six rdy_clr pulses rapidly consume the queue; count returns to zero and rdy falls.")

    doc.add_page_break()
    add_caption(doc, "Figure 3. Standalone FIFO capacity, simultaneous access, and wrap-around detail (TC10-TC12)")
    doc.add_picture(str(ASSETS / "uart_xsim_waveform_fifo_detail.png"), width=Inches(6.35))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_heading(doc, "5.3  FIFO unit behavior", 2)
    add_bullet(doc, "Sixteen write pulses increment count from zero to 16. fifo_full asserts exactly at the configured depth.")
    add_bullet(doc, "The extra overflow write does not increase count beyond 16.")
    add_bullet(doc, "Sixteen read pulses return count to zero and assert fifo_empty.")
    add_bullet(doc, "In TC11, read and write are high together; occupancy remains one and D2 becomes the new front value.")
    add_bullet(doc, "TC12 performs another write/read after the pointers have wrapped, demonstrating continued correct reuse.")

    add_heading(doc, "6. Conclusion", 1)
    add_body(
        doc,
        "The receiver-side FIFO is fully integrated without changing the top-level receive handshake. The design now buffers bursts on both sides of the UART, preserves byte order, handles delayed reads, flushes cleanly on reset, and recovers for later communication. Vivado XSim verified every planned basic case, and the saved WCFG/VCD artifacts make the result repeatable and inspectable."
    )
    add_callout(
        doc,
        "Result",
        "12/12 self-checking cases passed, zero reported errors, final simulation time 4098 ns.",
        fill=GREEN,
        title_color=GREEN_TEXT,
    )
    add_heading(doc, "6.1 Sensible next improvements", 2)
    add_bullet(doc, "Expose TX-full, RX-full, or available-count status to external logic.")
    add_bullet(doc, "Add a two-flop synchronizer for a real asynchronous RX pin.")
    add_bullet(doc, "Report framing error and optional parity error conditions.")
    add_bullet(doc, "Use majority voting around the center sample for improved noise tolerance.")
    add_bullet(doc, "Make data bits, parity, stop bits, FIFO depth, clock frequency, and baud rate configurable.")

    add_code_listing(doc, ROOT / "rtlCode" / "uart_top.v", "Appendix A - Integrated UART Top with TX and RX FIFOs")
    add_code_listing(doc, ROOT / "rtlCode" / "uart_fifo.v", "Appendix B - 16x8 Synchronous FIFO")
    add_code_listing(doc, ROOT / "rtlCode" / "uart_top_tb.v", "Appendix C - Self-checking Vivado Testbench")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_report()
